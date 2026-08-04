"""Generate DOCX version of the demo walkthrough."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- Title Page ---
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Multi-Database for AI', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('M&E Vertical POC — Complete Demo Walkthrough', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('4-Tier Architecture: DynamoDB → Aurora pgvector → Neptune Analytics → ElastiCache Valkey')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Target Customers: ').bold = True
p.add_run('Business Wire • Match Group • Universal Music Group • IMAX • Particle Media')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Pipeline: $2.7M+ ARR')

doc.add_page_break()

# --- Table of Contents ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Step 0: Launch the Application',
    'Step 1: Dashboard Overview',
    'Step 2: Event Ingestion',
    'Step 3: Semantic Analysis (pgvector)',
    'Step 4: Graph Intelligence (Neptune)',
    'Step 5: Real-Time Scoring (ElastiCache)',
    'Step 6: Full Pipeline Execution — The Main Event',
    'Step 7: AI Investigator Briefing (Claude)',
    'Step 8: Advanced Custom Scenarios',
    'Step 9: API Reference for Technical Audience',
    'Key Talking Points',
    'Quick Reference',
]
for i, item in enumerate(toc_items, 1):
    doc.add_paragraph(f'{i}. {item}', style='List Number')

doc.add_page_break()

# --- Step 0 ---
doc.add_heading('Step 0: Launch the Application', level=1)
doc.add_paragraph('Start the frontend development server:')
doc.add_paragraph('cd ~/DMS_local_converter/multi-db-poc/frontend\nnpm run dev', style='No Spacing')
doc.add_paragraph()
doc.add_paragraph('Open your browser to: http://localhost:5173')
doc.add_paragraph()
doc.add_paragraph('API endpoint: https://nkt0mgbdn5.execute-api.us-east-1.amazonaws.com/v1')
doc.add_paragraph()
doc.add_paragraph('The sidebar contains:')
doc.add_paragraph('Customer domain selector (dropdown) — switch between 5 customers', style='List Bullet')
doc.add_paragraph('Navigation links for all 7 pages', style='List Bullet')
doc.add_paragraph('Architecture tier indicators', style='List Bullet')

doc.add_page_break()

# --- Step 1 ---
doc.add_heading('Step 1: Dashboard Overview', level=1)
doc.add_paragraph('Click "Dashboard" in the sidebar navigation.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('What you see:').bold = True
doc.add_paragraph('5 metric cards: Events Ingested, Detections, Rings Discovered, Avg Latency, Cache Hit Rate', style='List Bullet')
doc.add_paragraph('Service Health panel: Green/yellow/red dots for DynamoDB, Aurora, Neptune, ElastiCache, Bedrock', style='List Bullet')
doc.add_paragraph('4-Tier Architecture diagram showing the data flow from ingestion to scoring', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Talking point: ').bold = True
p.add_run('"This dashboard gives us real-time visibility into all 4 database tiers. Everything is serverless — we pay only for what we use."')

doc.add_page_break()

# --- Step 2 ---
doc.add_heading('Step 2: Event Ingestion', level=1)
doc.add_paragraph('Click "Events" in the sidebar.')
doc.add_paragraph()
doc.add_heading('Example: Submit a Romance Scam Message (Match Group)', level=2)
doc.add_paragraph('1. Set domain to "Match Group" in the sidebar dropdown')
doc.add_paragraph('2. Fill in the form:')

table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'user_id'
table.rows[1].cells[1].text = 'USR-SUSPICIOUS-100'
table.rows[2].cells[0].text = 'recipient_id'
table.rows[2].cells[1].text = 'USR-REAL-005'
table.rows[3].cells[0].text = 'message_text'
table.rows[3].cells[1].text = 'I am deployed overseas'

doc.add_paragraph()
doc.add_paragraph('3. Click "Ingest Event"')
doc.add_paragraph('4. View the response in the right panel — you get back an event_id and timestamp')
doc.add_paragraph('5. Click "Refresh" to see your event in the Recent Events table')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Talking point: ').bold = True
p.add_run('"Every event is immediately captured in DynamoDB with sub-10ms latency. DynamoDB Streams then trigger the downstream analysis."')

doc.add_page_break()

# --- Step 3 ---
doc.add_heading('Step 3: Semantic Analysis (Aurora pgvector)', level=1)
doc.add_paragraph('Click "Semantic" in the sidebar.')
doc.add_paragraph()
doc.add_heading('Example: Detect a Romance Scam Script', level=2)
doc.add_paragraph('1. Set domain to "Match Group"')
doc.add_paragraph('2. In the textarea, paste this message:')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Hi sweetheart, I\'m a US Army captain deployed overseas. Looking for someone real to share my life with when I return home. You caught my eye. Can we talk on WhatsApp?').italic = True
doc.add_paragraph()
doc.add_paragraph('3. Set Similarity Threshold slider to 0.6')
doc.add_paragraph('4. Click "Run Similarity Search"')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Expected Results:').bold = True

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Match ID'
table.rows[0].cells[1].text = 'Score'
table.rows[0].cells[2].text = 'Interpretation'
table.rows[1].cells[0].text = 'SCAM-007'
table.rows[1].cells[1].text = '92.3%'
table.rows[1].cells[2].text = 'Near-identical to known scam (RED)'
table.rows[2].cells[0].text = 'SCAM-001'
table.rows[2].cells[1].text = '87.1%'
table.rows[2].cells[2].text = 'Strong similarity (RED)'
table.rows[3].cells[0].text = 'SCAM-005'
table.rows[3].cells[1].text = '79.5%'
table.rows[3].cells[2].text = 'Moderate match (YELLOW)'

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Talking point: ').bold = True
p.add_run('"This message was never seen before, but pgvector finds it\'s 92% semantically identical to a known scam script. Bedrock Titan V2 embeddings capture meaning — not just keywords. Even with different wording, the intent is detected."')

doc.add_paragraph()
doc.add_heading('False Positive Check: Test a Legitimate Message', level=2)
doc.add_paragraph('Paste this normal dating message:')
p = doc.add_paragraph()
p.add_run('Hey! I noticed we both like hiking and coffee. I work downtown as a designer. Would you want to grab a drink sometime this week?').italic = True
doc.add_paragraph()
doc.add_paragraph('Expected: 0 matches. Legitimate messages don\'t trigger the scam detection patterns.')

doc.add_page_break()

# --- Step 4 ---
doc.add_heading('Step 4: Graph Intelligence (Neptune Analytics)', level=1)
doc.add_paragraph('Click "Graph" in the sidebar.')
doc.add_paragraph()
doc.add_heading('Example: Discover the Romance Scam Ring', level=2)
doc.add_paragraph('1. Set domain to "Match Group"')
doc.add_paragraph('2. Enter Entity ID: USR-FAKE-001')
doc.add_paragraph('3. Algorithm: Louvain (community detection)')
doc.add_paragraph('4. Max Depth: 3')
doc.add_paragraph('5. Click "Run Graph Analysis"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Expected: ').bold = True
p.add_run('Neptune returns the full community — showing USR-FAKE-001 connected to USR-FAKE-004, USR-FAKE-007, USR-FAKE-010 via shared devices (DEV-SHARED-0, DEV-SHARED-1). Also shows contact with victim USR-REAL-003 (Amanda).')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Talking point: ').bold = True
p.add_run('"Neptune Analytics discovers that USR-FAKE-001 shares a device with multiple fake accounts. This is a coordinated 15-member ring sharing just 3 devices. A traditional relational database can\'t traverse these relationships in real-time — it would require expensive recursive JOINs."')
doc.add_paragraph()

doc.add_heading('Other Algorithms to Show', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Algorithm'
table.rows[0].cells[1].text = 'What it shows'
table.rows[0].cells[2].text = 'Best for'
table.rows[1].cells[0].text = 'Louvain'
table.rows[1].cells[1].text = 'Community clusters'
table.rows[1].cells[2].text = 'Finding rings/groups'
table.rows[2].cells[0].text = 'PageRank'
table.rows[2].cells[1].text = 'Node influence ranking'
table.rows[2].cells[2].text = 'Finding ring leaders'
table.rows[3].cells[0].text = 'WCC'
table.rows[3].cells[1].text = 'Connected components'
table.rows[3].cells[2].text = 'Full network mapping'
table.rows[4].cells[0].text = 'Shortest Path'
table.rows[4].cells[1].text = 'Path between entities'
table.rows[4].cells[2].text = 'Evidence chains'

doc.add_page_break()

# --- Step 5 ---
doc.add_heading('Step 5: Real-Time Scoring (ElastiCache Valkey)', level=1)
doc.add_paragraph('Click "Scoring" in the sidebar.')
doc.add_paragraph()
doc.add_heading('Example: Look Up a Known Scammer', level=2)
doc.add_paragraph('1. Enter Entity ID: USR-FAKE-001')
doc.add_paragraph('2. Click "Get Score"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Expected display:').bold = True
doc.add_paragraph('• Large "74" score gauge')
doc.add_paragraph('• CHALLENGE badge (orange)')
doc.add_paragraph('• Component bars: Graph 100%, Similarity 100%, Behavioral 30%, Velocity 20%')
doc.add_paragraph('• Cache indicator: "⚡ Cache Hit | 0.04ms"')
doc.add_paragraph()
doc.add_heading('Decision Matrix', level=2)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Score'
table.rows[0].cells[1].text = 'Decision'
table.rows[0].cells[2].text = 'Action'
table.rows[1].cells[0].text = '0 - 30'
table.rows[1].cells[1].text = 'ALLOW'
table.rows[1].cells[2].text = 'No action needed'
table.rows[2].cells[0].text = '30 - 60'
table.rows[2].cells[1].text = 'FLAG'
table.rows[2].cells[2].text = 'Manual review required'
table.rows[3].cells[0].text = '60 - 80'
table.rows[3].cells[1].text = 'CHALLENGE'
table.rows[3].cells[2].text = 'Identity verification required'
table.rows[4].cells[0].text = '80 - 100'
table.rows[4].cells[1].text = 'BLOCK'
table.rows[4].cells[2].text = 'Immediate block'

doc.add_page_break()

# --- Step 6 ---
doc.add_heading('Step 6: Full Pipeline Execution — The Main Event', level=1)
doc.add_paragraph('Click "Demo" in the sidebar. This is the most impressive demonstration.')
doc.add_paragraph()

# Match Group
doc.add_heading('Demo 1: Match Group — "The Romance Scam Ring"', level=2)
doc.add_paragraph('1. Set domain to "Match Group"')
doc.add_paragraph('2. Click "▶️ Execute Full Pipeline"')
doc.add_paragraph('3. Watch stages appear with results:')
doc.add_paragraph()

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Stage'
table.rows[0].cells[1].text = 'Status'
table.rows[0].cells[2].text = 'Latency'
table.rows[0].cells[3].text = 'Result'
table.rows[1].cells[0].text = 'Cache Check'
table.rows[1].cells[1].text = 'MISS'
table.rows[1].cells[2].text = '<1ms'
table.rows[1].cells[3].text = 'First time seeing this entity'
table.rows[2].cells[0].text = 'DynamoDB Ingest'
table.rows[2].cells[1].text = 'SUCCESS'
table.rows[2].cells[2].text = '8ms'
table.rows[2].cells[3].text = 'Event captured'
table.rows[3].cells[0].text = 'Bedrock Embedding'
table.rows[3].cells[1].text = 'SUCCESS'
table.rows[3].cells[2].text = '125ms'
table.rows[3].cells[3].text = '1024-dim vector generated'
table.rows[4].cells[0].text = 'pgvector Similarity'
table.rows[4].cells[1].text = 'SUCCESS'
table.rows[4].cells[2].text = '135ms'
table.rows[4].cells[3].text = '100% match to known scam script'
table.rows[5].cells[0].text = 'Neptune Graph'
table.rows[5].cells[1].text = 'SUCCESS'
table.rows[5].cells[2].text = '100ms'
table.rows[5].cells[3].text = 'Part of 15-member ring'
table.rows[6].cells[0].text = 'Composite Score'
table.rows[6].cells[1].text = 'SUCCESS'
table.rows[6].cells[2].text = '<1ms'
table.rows[6].cells[3].text = 'Score: 0.74 → CHALLENGE'

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Total: 370ms | Score: 74 | Decision: CHALLENGE').bold = True
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key message: ').bold = True
p.add_run('"In 370 milliseconds, across 4 purpose-built databases and an AI embedding model, we identified a known scam script, mapped the criminal ring, and issued a CHALLENGE decision. Legacy databases can\'t do vector similarity AND graph traversal in a single sub-second pipeline."')

doc.add_paragraph()

# Business Wire
doc.add_heading('Demo 2: Business Wire — "The Embargo Breach"', level=2)
doc.add_paragraph('Switch domain to "Business Wire" → Execute Pipeline')
doc.add_paragraph()
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Metric'
table.rows[0].cells[1].text = 'Result'
table.rows[1].cells[0].text = 'Similarity'
table.rows[1].cells[1].text = '100% match to breach access patterns'
table.rows[2].cells[0].text = 'Graph'
table.rows[2].cells[1].text = 'J-UNKNOWN-443 in 3-node leak network'
table.rows[3].cells[0].text = 'Decision'
table.rows[3].cells[1].text = 'CHALLENGE (0.74) — 377ms'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key message: ').bold = True
p.add_run('"In the real Business Wire hack (2010-2015), it took years to detect 150,000 stolen releases. This system catches it in 377 milliseconds."')

doc.add_paragraph()

# UMG
doc.add_heading('Demo 3: Universal Music — "The Stream Farm"', level=2)
doc.add_paragraph('Switch to "Universal Music" → Execute Pipeline')
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Metric'
table.rows[0].cells[1].text = 'Result'
table.rows[1].cells[0].text = 'Similarity'
table.rows[1].cells[1].text = '100% match to bot farm listening pattern'
table.rows[2].cells[0].text = 'Graph'
table.rows[2].cells[1].text = 'BOT-FARM-001 in 47-account network'
table.rows[3].cells[0].text = 'Decision'
table.rows[3].cells[1].text = 'CHALLENGE (0.63) — 340ms'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key message: ').bold = True
p.add_run('"Michael Smith pled guilty in March 2026 to $8M in streaming fraud. This catches the bot farm before a single fraudulent royalty is paid."')

doc.add_paragraph()

# IMAX
doc.add_heading('Demo 4: IMAX — "The Scalper Bot Network"', level=2)
doc.add_paragraph('Switch to "IMAX" → Execute Pipeline')
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Metric'
table.rows[0].cells[1].text = 'Result'
table.rows[1].cells[0].text = 'Similarity'
table.rows[1].cells[1].text = '100% match to bot behavioral pattern (85ms speed)'
table.rows[2].cells[0].text = 'Graph'
table.rows[2].cells[1].text = '23-session coordinated network'
table.rows[3].cells[0].text = 'Decision'
table.rows[3].cells[1].text = 'FLAG (0.55) — 374ms'

doc.add_paragraph()

# Particle Media
doc.add_heading('Demo 5: Particle Media — "The Misinformation Campaign"', level=2)
doc.add_paragraph('Switch to "Particle Media" → Execute Pipeline')
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Metric'
table.rows[0].cells[1].text = 'Result'
table.rows[1].cells[0].text = 'Similarity'
table.rows[1].cells[1].text = '100% match to AI-generated misinfo patterns'
table.rows[2].cells[0].text = 'Graph'
table.rows[2].cells[1].text = '50-account amplification network'
table.rows[3].cells[0].text = 'Decision'
table.rows[3].cells[1].text = 'CHALLENGE (0.74) — 379ms'

doc.add_page_break()

# --- Step 7 ---
doc.add_heading('Step 7: AI Investigator Briefing (Claude via Bedrock)', level=1)
doc.add_paragraph('Click "Briefing" in the sidebar.')
doc.add_paragraph()
doc.add_paragraph('1. Enter Entity ID: USR-FAKE-001')
doc.add_paragraph('2. Click "Generate Briefing"')
doc.add_paragraph('3. Wait 3-5 seconds')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Claude generates a full investigator-grade report containing:').bold = True
doc.add_paragraph()
doc.add_paragraph('Executive Summary — 2-3 sentence overview of the threat', style='List Bullet')
doc.add_paragraph('Entity Profile — Classification, ring position, connected accounts, infrastructure', style='List Bullet')
doc.add_paragraph('Evidence Timeline — Sequenced events with severity ratings', style='List Bullet')
doc.add_paragraph('Risk Assessment — Critical/High/Medium/Low with detailed justification', style='List Bullet')
doc.add_paragraph('Recommended Actions — 8 prioritized items with owners and timelines:', style='List Bullet')
doc.add_paragraph('  1. Immediate account suspension (2 hours)', style='List Bullet')
doc.add_paragraph('  2. Victim notification (1 hour)', style='List Bullet')
doc.add_paragraph('  3. Device forensics (24 hours)', style='List Bullet')
doc.add_paragraph('  4. Law enforcement escalation — FBI IC3', style='List Bullet')
doc.add_paragraph('  5. Ring expansion investigation', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Talking point: ').bold = True
p.add_run('"Claude synthesizes evidence from all 4 tiers into a production-quality investigation report. An analyst who would spend hours writing this gets it in 5 seconds."')

doc.add_page_break()

# --- Key Talking Points ---
doc.add_heading('Key Talking Points Summary', level=1)
doc.add_paragraph()

table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Point'
table.rows[0].cells[1].text = 'Evidence'
table.rows[1].cells[0].text = 'Sub-400ms detection'
table.rows[1].cells[1].text = 'Pipeline stages total 340-380ms warm'
table.rows[2].cells[0].text = '4 purpose-built databases'
table.rows[2].cells[1].text = 'Each tier handles what it\'s best at'
table.rows[3].cells[0].text = 'AI-powered semantic understanding'
table.rows[3].cells[1].text = 'Bedrock Titan V2 catches meaning, not keywords'
table.rows[4].cells[0].text = 'Graph intelligence finds hidden networks'
table.rows[4].cells[1].text = 'Neptune maps rings invisible to flat queries'
table.rows[5].cells[0].text = 'One codebase, 5 customers'
table.rows[5].cells[1].text = 'Domain selector swaps everything'
table.rows[6].cells[0].text = 'Serverless, pay-per-use'
table.rows[6].cells[1].text = 'DynamoDB on-demand, Aurora scales, Lambda per-invoke'
table.rows[7].cells[0].text = 'Investigation reports in seconds'
table.rows[7].cells[1].text = 'Claude generates analyst-grade briefings'

doc.add_paragraph()

# --- Quick Reference ---
doc.add_heading('Quick Reference: Entity IDs', level=1)
doc.add_paragraph()

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Customer'
table.rows[0].cells[1].text = 'Entity ID'
table.rows[0].cells[2].text = 'Scenario'
table.rows[1].cells[0].text = 'Match Group'
table.rows[1].cells[1].text = 'USR-FAKE-001'
table.rows[1].cells[2].text = 'Romance scam ring member'
table.rows[2].cells[0].text = 'Business Wire'
table.rows[2].cells[1].text = 'J-UNKNOWN-443'
table.rows[2].cells[2].text = 'Unauthorized journalist'
table.rows[3].cells[0].text = 'Universal Music'
table.rows[3].cells[1].text = 'BOT-FARM-001'
table.rows[3].cells[2].text = 'Stream farm bot account'
table.rows[4].cells[0].text = 'IMAX'
table.rows[4].cells[1].text = 'SESS-BOT-001'
table.rows[4].cells[2].text = 'Scalper bot session'
table.rows[5].cells[0].text = 'Particle Media'
table.rows[5].cells[1].text = 'PM-BOT-001'
table.rows[5].cells[2].text = 'Misinfo amplification bot'

doc.add_paragraph()
doc.add_heading('Demo Timing', level=1)

table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Section'
table.rows[0].cells[1].text = 'Time'
table.rows[0].cells[2].text = 'Priority'
table.rows[1].cells[0].text = 'Dashboard overview'
table.rows[1].cells[1].text = '1 min'
table.rows[1].cells[2].text = 'Must'
table.rows[2].cells[0].text = 'Pipeline: Match Group'
table.rows[2].cells[1].text = '3 min'
table.rows[2].cells[2].text = 'Must'
table.rows[3].cells[0].text = 'Pipeline: Business Wire'
table.rows[3].cells[1].text = '2 min'
table.rows[3].cells[2].text = 'Must'
table.rows[4].cells[0].text = 'Pipeline: 1-2 more domains'
table.rows[4].cells[1].text = '3 min'
table.rows[4].cells[2].text = 'Should'
table.rows[5].cells[0].text = 'Semantic Analysis deep-dive'
table.rows[5].cells[1].text = '2 min'
table.rows[5].cells[2].text = 'Should'
table.rows[6].cells[0].text = 'Claude Briefing generation'
table.rows[6].cells[1].text = '2 min'
table.rows[6].cells[2].text = 'Must (wow factor)'
table.rows[7].cells[0].text = 'Graph Intelligence'
table.rows[7].cells[1].text = '2 min'
table.rows[7].cells[2].text = 'If time allows'
table.rows[8].cells[0].text = 'TOTAL'
table.rows[8].cells[1].text = '~15 min'
table.rows[8].cells[2].text = ''

# Save
output_path = '/Users/haliasgh/DMS_local_converter/multi-db-poc/docs/Multi-DB_AI_POC_Demo_Walkthrough.docx'
doc.save(output_path)
print(f'Generated: {output_path}')
