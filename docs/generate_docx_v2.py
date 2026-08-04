"""Generate comprehensive DOCX demo walkthrough with full data for all 5 demos."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)


def add_talking_point(doc, text):
    p = doc.add_paragraph()
    p.add_run('💬 Talking Point: ').bold = True
    run = p.add_run(f'"{text}"')
    run.italic = True


def add_stage_table(doc, stages):
    table = doc.add_table(rows=len(stages)+1, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = ['Stage', 'Status', 'Latency', 'Result']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_idx, stage in enumerate(stages, 1):
        for col_idx, val in enumerate(stage):
            table.rows[row_idx].cells[col_idx].text = val


# ===== TITLE PAGE =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Multi-Database for AI: M&E Vertical POC', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Complete Demo Walkthrough Guide', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DynamoDB → Aurora pgvector → Neptune Analytics → ElastiCache Valkey')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('5 Customers: ').bold = True
p.add_run('Business Wire • Match Group • Universal Music • IMAX • Particle Media')
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Total Pipeline: $2.7M+ ARR | Performance Target: <540ms | Achieved: <400ms')
doc.add_page_break()

# ===== TABLE OF CONTENTS =====
doc.add_heading('Table of Contents', level=1)
sections = [
    'Step 0: Launch the Application',
    'Step 1: Dashboard Overview',
    'Step 2: Event Ingestion — Submitting Events',
    'Step 3: Semantic Analysis — Vector Similarity Detection',
    'Step 4: Graph Intelligence — Network Discovery',
    'Step 5: Real-Time Scoring — Risk Assessment',
    'Step 6: Full Pipeline Demo 1 — Match Group (Romance Scam Ring)',
    'Step 7: Full Pipeline Demo 2 — Business Wire (Embargo Breach)',
    'Step 8: Full Pipeline Demo 3 — Universal Music (Stream Farm)',
    'Step 9: Full Pipeline Demo 4 — IMAX (Scalper Bot Network)',
    'Step 10: Full Pipeline Demo 5 — Particle Media (Misinformation Campaign)',
    'Step 11: AI Investigator Briefing (Claude)',
    'Step 12: Advanced — Custom Scenarios & False Positive Testing',
    'Step 13: API Reference for Technical Audience',
    'Appendix A: Entity IDs Quick Reference',
    'Appendix B: Architecture & Decision Matrix',
    'Appendix C: Demo Timing Guide',
]
for i, s in enumerate(sections, 1):
    doc.add_paragraph(f'{i}. {s}', style='List Number')
doc.add_page_break()

# ===== STEP 0 =====
doc.add_heading('Step 0: Launch the Application', level=1)
doc.add_paragraph('Prerequisites:')
doc.add_paragraph('Node.js 18+ installed', style='List Bullet')
doc.add_paragraph('AWS credentials active (run mwinit if expired)', style='List Bullet')
doc.add_paragraph('Network access to API Gateway from your IP', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Start the frontend:')
add_code_block(doc, 'cd ~/DMS_local_converter/multi-db-poc/frontend\nnpm run dev')
doc.add_paragraph()
doc.add_paragraph('Open browser: http://localhost:5173')
doc.add_paragraph()
doc.add_paragraph('Verify API is healthy:')
add_code_block(doc, 'curl https://nkt0mgbdn5.execute-api.us-east-1.amazonaws.com/v1/health\n# Expected: {"status":"ok"}')
doc.add_page_break()

# ===== STEP 1 =====
doc.add_heading('Step 1: Dashboard Overview', level=1)
doc.add_paragraph('Navigate: Click "📊 Dashboard" in the sidebar')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Elements on screen:').bold = True
doc.add_paragraph('5 Metric Cards — Events Ingested, Detections, Rings Discovered, Avg Latency, Cache Hit Rate', style='List Bullet')
doc.add_paragraph('Service Health — Green/yellow/red status for each of 5 services', style='List Bullet')
doc.add_paragraph('Architecture Diagram — 4 color-coded tiers showing the pipeline flow', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Domain Selector: Use the dropdown at the top of the sidebar to switch between customers. All pages adapt to the selected domain.')
doc.add_paragraph()
add_talking_point(doc, 'This dashboard gives real-time visibility into all 4 database tiers. Everything is serverless — DynamoDB on-demand, Aurora scales between 2-8 ACUs, Lambda per-invocation. We pay only for what we use.')
doc.add_page_break()

# ===== STEP 2 =====
doc.add_heading('Step 2: Event Ingestion — Submitting Events', level=1)
doc.add_paragraph('Navigate: Click "📥 Events" in the sidebar')
doc.add_paragraph()
doc.add_paragraph('The form fields adapt per customer domain. Here are examples for each:')
doc.add_paragraph()

doc.add_heading('Business Wire — Embargo Access Event', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'release_id'
table.rows[1].cells[1].text = 'PR-2024-0004'
table.rows[2].cells[0].text = 'journalist_id'
table.rows[2].cells[1].text = 'J-UNAUTHORIZED-99'
table.rows[3].cells[0].text = 'access_type'
table.rows[3].cells[1].text = 'unauthorized_early_access'
table.rows[4].cells[0].text = 'content'
table.rows[4].cells[1].text = 'MegaCorp acquisition of TechStartup accessed 4 hours before embargo lift'
doc.add_paragraph()

doc.add_heading('Match Group — Scam Message Event', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'user_id'
table.rows[1].cells[1].text = 'USR-FAKE-001'
table.rows[2].cells[0].text = 'recipient_id'
table.rows[2].cells[1].text = 'USR-REAL-005'
table.rows[3].cells[0].text = 'message_text'
table.rows[3].cells[1].text = 'Hello beautiful, I am a military officer overseas'
doc.add_paragraph()

doc.add_heading('Universal Music — Stream Event', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'account_id'
table.rows[1].cells[1].text = 'BOT-FARM-001'
table.rows[2].cells[0].text = 'track_id'
table.rows[2].cells[1].text = 'AI-TRACK-001'
table.rows[3].cells[0].text = 'artist'
table.rows[3].cells[1].text = 'AI_Artist_1'
table.rows[4].cells[0].text = 'duration_ms'
table.rows[4].cells[1].text = '300 (0.3 seconds — bot signature)'
doc.add_paragraph()

doc.add_heading('IMAX — Purchase Attempt Event', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'session_id'
table.rows[1].cells[1].text = 'SESS-BOT-001'
table.rows[2].cells[0].text = 'showtime_id'
table.rows[2].cells[1].text = 'IMAX-PREM-001'
table.rows[3].cells[0].text = 'quantity'
table.rows[3].cells[1].text = '8'
table.rows[4].cells[0].text = 'device_fingerprint'
table.rows[4].cells[1].text = 'BOT-FP-0'
doc.add_paragraph()

doc.add_heading('Particle Media — Content Published Event', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'content_id'
table.rows[1].cells[1].text = 'MISINFO-2024-001'
table.rows[2].cells[0].text = 'author_id'
table.rows[2].cells[1].text = 'BOT-AUTHOR-50'
table.rows[3].cells[0].text = 'content_text'
table.rows[3].cells[1].text = 'BREAKING: Vaccine causes 90% side effects'
table.rows[4].cells[0].text = 'source_url'
table.rows[4].cells[1].text = 'fake-news-site.com'
doc.add_paragraph()
doc.add_paragraph('Expected response for all: event_id (UUID), timestamp, status: "ingested", latency: <10ms')
doc.add_page_break()

# ===== STEP 3 =====
doc.add_heading('Step 3: Semantic Analysis — Vector Similarity Detection', level=1)
doc.add_paragraph('Navigate: Click "🧠 Semantic" in the sidebar')
doc.add_paragraph()
doc.add_paragraph('This page runs Bedrock Titan V2 embeddings (1024 dimensions) and cosine similarity search against known threat patterns stored in Aurora pgvector.')
doc.add_paragraph()

doc.add_heading('Example 1: Match Group — Detect Romance Scam', level=2)
doc.add_paragraph('Domain: Match Group | Threshold: 0.6')
doc.add_paragraph('Input text:')
add_code_block(doc, 'Hi sweetheart, I\'m a US Army captain deployed overseas. Looking for\nsomeone real to share my life with when I return home. You caught\nmy eye. Can we talk on WhatsApp?')
doc.add_paragraph()
doc.add_paragraph('Expected matches:')
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Match ID'
table.rows[0].cells[1].text = 'Cosine Score'
table.rows[0].cells[2].text = 'Known Pattern'
table.rows[1].cells[0].text = 'SCAM-007'
table.rows[1].cells[1].text = '92.3%'
table.rows[1].cells[2].text = 'US Army captain deployed, looking for someone real'
table.rows[2].cells[0].text = 'SCAM-001'
table.rows[2].cells[1].text = '87.1%'
table.rows[2].cells[2].text = 'US military officer, move to WhatsApp'
table.rows[3].cells[0].text = 'SCAM-005'
table.rows[3].cells[1].text = '79.5%'
table.rows[3].cells[2].text = 'NATO forces deployed, limited internet'
doc.add_paragraph()
add_talking_point(doc, 'pgvector found this is 92% semantically identical to known scam scripts. Titan V2 captures meaning — even with different wording, the romance scam intent is detected.')

doc.add_paragraph()
doc.add_heading('Example 2: Business Wire — Detect Embargo Breach Pattern', level=2)
doc.add_paragraph('Domain: Business Wire | Threshold: 0.6')
doc.add_paragraph('Input text:')
add_code_block(doc, 'Unauthorized access detected: unverified journalist account accessed\nembargoed M&A announcement for MegaCorp acquisition 4 hours\nbefore scheduled publication time from suspicious Eastern European IP.')
doc.add_paragraph()
doc.add_paragraph('Expected matches:')
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Match ID'
table.rows[0].cells[1].text = 'Score'
table.rows[0].cells[2].text = 'Pattern'
table.rows[1].cells[0].text = 'BREACH-003'
table.rows[1].cells[1].text = '89%'
table.rows[1].cells[2].text = 'Coordinated breach, Eastern European IP, embargoed financial releases'
table.rows[2].cells[0].text = 'BREACH-001'
table.rows[2].cells[1].text = '84%'
table.rows[2].cells[2].text = 'Unauthorized early access from unverified journalist'
table.rows[3].cells[0].text = 'BREACH-002'
table.rows[3].cells[1].text = '76%'
table.rows[3].cells[2].text = 'Same IP accessed multiple embargoed releases'
doc.add_paragraph()

doc.add_heading('Example 3: UMG — Detect Bot Farm Listening Pattern', level=2)
doc.add_paragraph('Domain: Universal Music | Threshold: 0.5')
doc.add_paragraph('Input text:')
add_code_block(doc, 'Account streaming over 500000 times per day with average listen\nduration of 2 seconds from a device fingerprint shared with 30\nother accounts all streaming the same AI-generated tracks exclusively.')
doc.add_paragraph()
doc.add_paragraph('Expected matches:')
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Match ID'
table.rows[0].cells[1].text = 'Score'
table.rows[0].cells[2].text = 'Pattern'
table.rows[1].cells[0].text = 'BOT-PAT-003'
table.rows[1].cells[1].text = '95%+'
table.rows[1].cells[2].text = '661K streams/day, 0.3s duration, shared device, 47-account network'
table.rows[2].cells[0].text = 'BOT-PAT-001'
table.rows[2].cells[1].text = '63%'
table.rows[2].cells[2].text = '50K+ streams/day, <5s duration, shared device fingerprint'
doc.add_paragraph()

doc.add_heading('Example 4: IMAX — Detect Scalper Bot Behavior', level=2)
doc.add_paragraph('Domain: IMAX | Threshold: 0.5')
doc.add_paragraph('Input text:')
add_code_block(doc, 'Automated ticket purchase detected: 150 sessions initiated within\n20 seconds for the same premium showing. All sessions have\ninteraction speed under 100ms and share 4 device fingerprints.')
doc.add_paragraph()
doc.add_paragraph('Expected matches:')
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Match ID'
table.rows[0].cells[1].text = 'Score'
table.rows[0].cells[2].text = 'Pattern'
table.rows[1].cells[0].text = 'SCALP-002'
table.rows[1].cells[1].text = '91%'
table.rows[1].cells[2].text = '200 sessions in 30s, shared fingerprints and payment BINs'
table.rows[2].cells[0].text = 'SCALP-001'
table.rows[2].cells[1].text = '78%'
table.rows[2].cells[2].text = 'Interaction speed 85ms, shared device, bulk quantity'
doc.add_paragraph()

doc.add_heading('Example 5: Particle Media — Detect AI Misinformation', level=2)
doc.add_paragraph('Domain: Particle Media | Threshold: 0.5')
doc.add_paragraph('Input text:')
add_code_block(doc, 'SHOCKING REPORT: Government study reveals 5G radiation causes\ncancer in 85% of nearby residents. Internal documents prove\na decade-long cover-up by major telecom corporations.')
doc.add_paragraph()
doc.add_paragraph('Expected matches:')
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Match ID'
table.rows[0].cells[1].text = 'Score'
table.rows[0].cells[2].text = 'Pattern'
table.rows[1].cells[0].text = 'MISINFO-001'
table.rows[1].cells[1].text = '88%'
table.rows[1].cells[2].text = 'AI-generated sensationalist health misinfo, cover-up claims'
table.rows[2].cells[0].text = 'MISINFO-003'
table.rows[2].cells[1].text = '82%'
table.rows[2].cells[2].text = 'Vaccine/health misinfo with leaked documents framing'
doc.add_paragraph()

doc.add_heading('False Positive Test (Important for Credibility)', level=2)
doc.add_paragraph('Domain: Match Group | Threshold: 0.6')
doc.add_paragraph('Input (legitimate message):')
add_code_block(doc, 'Hey! I noticed we both like hiking and coffee. I work downtown\nas a designer. Would you want to grab a drink sometime this week?')
doc.add_paragraph('Expected: 0 matches. System correctly identifies normal messages as non-threatening.')
doc.add_page_break()

# ===== STEP 4 =====
doc.add_heading('Step 4: Graph Intelligence — Network Discovery', level=1)
doc.add_paragraph('Navigate: Click "🕸️ Graph" in the sidebar')
doc.add_paragraph()

doc.add_heading('Example 1: Match Group — Map the 15-Member Scam Ring', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Setting'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Entity ID'
table.rows[1].cells[1].text = 'USR-FAKE-001'
table.rows[2].cells[0].text = 'Algorithm'
table.rows[2].cells[1].text = 'Louvain (community detection)'
table.rows[3].cells[0].text = 'Max Depth'
table.rows[3].cells[1].text = '3'
doc.add_paragraph()
doc.add_paragraph('Expected results — nodes discovered:')
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Node ID'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Properties'
table.rows[1].cells[0].text = 'USR-FAKE-001'
table.rows[1].cells[1].text = 'User (scam)'
table.rows[1].cells[2].text = 'device: DEV-SHARED-0'
table.rows[2].cells[0].text = 'USR-FAKE-004'
table.rows[2].cells[1].text = 'User (scam)'
table.rows[2].cells[2].text = 'device: DEV-SHARED-1'
table.rows[3].cells[0].text = 'USR-FAKE-007'
table.rows[3].cells[1].text = 'User (scam)'
table.rows[3].cells[2].text = 'device: DEV-SHARED-1'
table.rows[4].cells[0].text = 'DEV-SHARED-0'
table.rows[4].cells[1].text = 'Device'
table.rows[4].cells[2].text = 'shared_device'
table.rows[5].cells[0].text = 'DEV-SHARED-1'
table.rows[5].cells[1].text = 'Device'
table.rows[5].cells[2].text = 'shared_device'
table.rows[6].cells[0].text = 'USR-REAL-003'
table.rows[6].cells[1].text = 'User (victim)'
table.rows[6].cells[2].text = 'name: Amanda'
doc.add_paragraph()
add_talking_point(doc, 'Neptune discovers that 15 fake accounts share just 3 devices. This is a coordinated ring. A SQL database would need expensive recursive JOINs — Neptune traverses this in 100ms.')

doc.add_paragraph()
doc.add_heading('Example 2: Business Wire — Map the Leak Network', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Setting'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Entity ID'
table.rows[1].cells[1].text = 'J-UNKNOWN-443'
table.rows[2].cells[0].text = 'Algorithm'
table.rows[2].cells[1].text = 'Louvain'
table.rows[3].cells[0].text = 'Depth'
table.rows[3].cells[1].text = '3'
doc.add_paragraph()
doc.add_paragraph('Expected: J-UNKNOWN-443 connected to J-006 and J-007 via SHARES_IP edges, all three connected to embargoed releases PR-2024-0001, PR-2024-0002, PR-2024-0003 via ACCESSED_EMBARGO edges.')
doc.add_paragraph()

doc.add_heading('Example 3: UMG — Map the 47-Account Bot Farm', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Setting'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Entity ID'
table.rows[1].cells[1].text = 'BOT-FARM-001'
table.rows[2].cells[0].text = 'Algorithm'
table.rows[2].cells[1].text = 'Louvain'
table.rows[3].cells[0].text = 'Depth'
table.rows[3].cells[1].text = '3'
doc.add_paragraph()
doc.add_paragraph('Expected: BOT-FARM-001 shares device BOT-DEV-0 with ~9 other bot accounts. All connected to AI-generated tracks via STREAMED edges. 47 accounts total in the connected component, sharing 5 devices.')
doc.add_paragraph()

doc.add_heading('Example 4: IMAX — Map the 23-Device Scalper Network', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Setting'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Entity ID'
table.rows[1].cells[1].text = 'SESS-BOT-001'
table.rows[2].cells[0].text = 'Algorithm'
table.rows[2].cells[1].text = 'Louvain'
table.rows[3].cells[0].text = 'Depth'
table.rows[3].cells[1].text = '3'
doc.add_paragraph()
doc.add_paragraph('Expected: SESS-BOT-001 shares device fingerprint BOT-FP-0 with ~4 other sessions. All 23 sessions target the same showtime IMAX-PREM-001 via TARGETED_SHOWTIME edges. Network uses 5 device fingerprints total.')
doc.add_paragraph()

doc.add_heading('Example 5: Particle Media — Map the 50-Account Amplification Network', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Setting'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'Entity ID'
table.rows[1].cells[1].text = 'PM-BOT-001'
table.rows[2].cells[0].text = 'Algorithm'
table.rows[2].cells[1].text = 'Louvain'
table.rows[3].cells[0].text = 'Depth'
table.rows[3].cells[1].text = '3'
doc.add_paragraph()
doc.add_paragraph('Expected: PM-BOT-001 is MEMBER_OF "AMPLIFICATION-NET-01" sharing network with 19 other bot accounts. Each connected to misinformation articles MISINFO-001 through MISINFO-005 via AMPLIFIED edges.')
doc.add_page_break()

# ===== STEP 5 =====
doc.add_heading('Step 5: Real-Time Scoring — Risk Assessment', level=1)
doc.add_paragraph('Navigate: Click "⚡ Scoring" in the sidebar')
doc.add_paragraph()
doc.add_paragraph('Look up any entity that has been through the pipeline:')
doc.add_paragraph()
table = doc.add_table(rows=6, cols=5)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Entity ID'
table.rows[0].cells[1].text = 'Composite'
table.rows[0].cells[2].text = 'Decision'
table.rows[0].cells[3].text = 'Sim'
table.rows[0].cells[4].text = 'Graph'
table.rows[1].cells[0].text = 'USR-FAKE-001'
table.rows[1].cells[1].text = '0.74'
table.rows[1].cells[2].text = 'CHALLENGE'
table.rows[1].cells[3].text = '100%'
table.rows[1].cells[4].text = '100%'
table.rows[2].cells[0].text = 'J-UNKNOWN-443'
table.rows[2].cells[1].text = '0.74'
table.rows[2].cells[2].text = 'CHALLENGE'
table.rows[2].cells[3].text = '100%'
table.rows[2].cells[4].text = '100%'
table.rows[3].cells[0].text = 'BOT-FARM-001'
table.rows[3].cells[1].text = '0.63'
table.rows[3].cells[2].text = 'CHALLENGE'
table.rows[3].cells[3].text = '100%'
table.rows[3].cells[4].text = '100%'
table.rows[4].cells[0].text = 'SESS-BOT-001'
table.rows[4].cells[1].text = '0.55'
table.rows[4].cells[2].text = 'FLAG'
table.rows[4].cells[3].text = '100%'
table.rows[4].cells[4].text = '100%'
table.rows[5].cells[0].text = 'PM-BOT-001'
table.rows[5].cells[1].text = '0.74'
table.rows[5].cells[2].text = 'CHALLENGE'
table.rows[5].cells[3].text = '100%'
table.rows[5].cells[4].text = '100%'
doc.add_page_break()

# ===== STEP 6: DEMO 1 - Match Group =====
doc.add_heading('Step 6: Full Pipeline Demo 1 — Match Group', level=1)
doc.add_heading('"The Romance Scam Ring"', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Real-world context: ').bold = True
p.add_run('Match Group reported AI deepfake breaches in 2026. Reality Defender collaboration announced. $1.1B in deepfake losses in 2025. 11-company anti-fraud accord signed.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Scenario: ').bold = True
p.add_run('A known scam ring member (USR-FAKE-001) sends a message using a scripted romance scam template. The system detects the scripted content, maps the criminal ring, and issues a CHALLENGE decision.')
doc.add_paragraph()
doc.add_paragraph('Navigate: Click "▶️ Demo" → Ensure "Match Group" is selected → Click "Execute Full Pipeline"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Input data sent:').bold = True
doc.add_paragraph()
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Parameter'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'domain'
table.rows[1].cells[1].text = 'match_group'
table.rows[2].cells[0].text = 'event_type'
table.rows[2].cells[1].text = 'message_sent'
table.rows[3].cells[0].text = 'entity_id'
table.rows[3].cells[1].text = 'USR-FAKE-001'
table.rows[4].cells[0].text = 'content'
table.rows[4].cells[1].text = 'Hello beautiful, I am a US military officer stationed overseas. I would love to get to know you better. Can we move to WhatsApp for more private conversation?'
table.rows[5].cells[0].text = 'payload.user_id'
table.rows[5].cells[1].text = 'USR-FAKE-001'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pipeline stages and results:').bold = True
doc.add_paragraph()
add_stage_table(doc, [
    ['Cache Check', 'MISS', '<1ms', 'No prior score cached — first analysis'],
    ['DynamoDB Ingest', 'SUCCESS', '8ms', 'Event stored in multidb_poc_user_interaction_events table'],
    ['Bedrock Embedding', 'SUCCESS', '125ms', '1024-dimension vector generated via Titan Embed V2'],
    ['pgvector Similarity', 'SUCCESS', '135ms', 'MATCH: SCAM-001 at 100% cosine similarity'],
    ['Neptune Graph', 'SUCCESS', '100ms', 'Ring detected: 15+ connected entities via shared devices'],
    ['Composite Score', 'SUCCESS', '<1ms', 'Score: 0.74 | Decision: CHALLENGE'],
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Final result:').bold = True
doc.add_paragraph('• Total latency: 370ms')
doc.add_paragraph('• Risk score: 74/100')
doc.add_paragraph('• Decision: CHALLENGE (identity verification required)')
doc.add_paragraph('• Similarity score: 1.00 (100% match to known scam script)')
doc.add_paragraph('• Graph score: 1.00 (member of large coordinated ring)')
doc.add_paragraph()
add_talking_point(doc, 'In 370 milliseconds, across 4 purpose-built databases and an AI model, we identified a known scam script, mapped a 15-member criminal ring sharing 3 devices, and issued a CHALLENGE decision requiring identity verification. Legacy databases cannot do vector similarity AND graph traversal in a single sub-second pipeline.')
doc.add_page_break()

# ===== STEP 7: DEMO 2 - Business Wire =====
doc.add_heading('Step 7: Full Pipeline Demo 2 — Business Wire', level=1)
doc.add_heading('"The Embargo Breach"', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Real-world context: ').bold = True
p.add_run('The 2010-2015 Business Wire hack resulted in $100M insider trading profits. 150,000 press releases stolen. 32 defendants charged by SEC/DOJ. Took 5 years to fully detect and prosecute.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Scenario: ').bold = True
p.add_run('An unknown actor (J-UNKNOWN-443) accesses a confidential M&A press release about MegaCorp acquiring TechStartup for $2.3B, 4 hours before the embargo lifts. The system detects the breach pattern, maps the leak network, and flags the activity.')
doc.add_paragraph()
doc.add_paragraph('Navigate: Select "Business Wire" → Demo page → Execute')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Input data sent:').bold = True
doc.add_paragraph()
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Parameter'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'domain'
table.rows[1].cells[1].text = 'business_wire'
table.rows[2].cells[0].text = 'event_type'
table.rows[2].cells[1].text = 'embargo_access'
table.rows[3].cells[0].text = 'entity_id'
table.rows[3].cells[1].text = 'J-UNKNOWN-443'
table.rows[4].cells[0].text = 'content'
table.rows[4].cells[1].text = 'CONFIDENTIAL: MegaCorp to acquire TechStartup for 2.3 billion dollars. Deal expected to close Q4. This is embargoed information not for distribution.'
table.rows[5].cells[0].text = 'payload.release_id'
table.rows[5].cells[1].text = 'PR-2024-0004'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pipeline stages and results:').bold = True
doc.add_paragraph()
add_stage_table(doc, [
    ['Cache Check', 'MISS', '<1ms', 'First time seeing this journalist entity'],
    ['DynamoDB Ingest', 'SUCCESS', '7ms', 'Stored in multidb_poc_press_release_events'],
    ['Bedrock Embedding', 'SUCCESS', '130ms', '1024-dim vector of M&A content generated'],
    ['pgvector Similarity', 'SUCCESS', '140ms', 'MATCH: BREACH-004 at 100% (identical M&A language)'],
    ['Neptune Graph', 'SUCCESS', '95ms', 'J-UNKNOWN-443 in 3-node leak network (shared IP with J-006, J-007)'],
    ['Composite Score', 'SUCCESS', '<1ms', 'Score: 0.74 | Decision: CHALLENGE'],
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Final result:').bold = True
doc.add_paragraph('• Total latency: 377ms')
doc.add_paragraph('• Risk score: 74/100')
doc.add_paragraph('• Decision: CHALLENGE')
doc.add_paragraph('• Similarity: 100% match to known embargo breach content pattern')
doc.add_paragraph('• Graph: Connected to 2 other unauthorized actors via shared IP 198.51.100.42')
doc.add_paragraph()
doc.add_paragraph('Graph network discovered:')
doc.add_paragraph('• J-UNKNOWN-443 → SHARES_IP → J-006', style='List Bullet')
doc.add_paragraph('• J-UNKNOWN-443 → SHARES_IP → J-007', style='List Bullet')
doc.add_paragraph('• All three → ACCESSED_EMBARGO → PR-2024-0001, PR-2024-0002, PR-2024-0003', style='List Bullet')
doc.add_paragraph()
add_talking_point(doc, 'In the real case, it took 5 years and a multi-agency investigation to uncover the $100M insider trading ring. This system detects the breach in 377 milliseconds and immediately maps the 3-person leak network via shared IP analysis in Neptune.')
doc.add_page_break()

# ===== STEP 8: DEMO 3 - UMG =====
doc.add_heading('Step 8: Full Pipeline Demo 3 — Universal Music Group', level=1)
doc.add_heading('"The Stream Farm ($10M Fraud Ring)"', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Real-world context: ').bold = True
p.add_run('Michael Smith pled guilty March 2026 to $8M streaming fraud, forfeited earnings. Apple Music demonetized 2 billion bot streams in 2025. Forbes estimated the "streaming fraud machine" at $4B industry-wide.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Scenario: ').bold = True
p.add_run('Bot account BOT-FARM-001 streams AI-generated tracks 661,000 times per day with 0.3-second average listen duration. Device fingerprint shared with 47 other accounts. System detects bot pattern, maps the full farm, and prevents fraudulent royalty payouts.')
doc.add_paragraph()
doc.add_paragraph('Navigate: Select "Universal Music" → Demo page → Execute')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Input data sent:').bold = True
doc.add_paragraph()
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Parameter'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'domain'
table.rows[1].cells[1].text = 'umg'
table.rows[2].cells[0].text = 'event_type'
table.rows[2].cells[1].text = 'stream'
table.rows[3].cells[0].text = 'entity_id'
table.rows[3].cells[1].text = 'BOT-FARM-001'
table.rows[4].cells[0].text = 'content'
table.rows[4].cells[1].text = 'Bot farm pattern detected: 661000 streams per day from single account, 0.3 second average duration, same device ID shared across 47 accounts in network.'
table.rows[5].cells[0].text = 'payload.account_id'
table.rows[5].cells[1].text = 'BOT-FARM-001'
table.rows[6].cells[0].text = 'payload.streams_per_day'
table.rows[6].cells[1].text = '661000'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pipeline stages and results:').bold = True
doc.add_paragraph()
add_stage_table(doc, [
    ['Cache Check', 'MISS', '<1ms', 'No cached score'],
    ['DynamoDB Ingest', 'SUCCESS', '9ms', 'Stored in multidb_poc_stream_events'],
    ['Bedrock Embedding', 'SUCCESS', '120ms', '1024-dim vector of bot behavior description'],
    ['pgvector Similarity', 'SUCCESS', '145ms', 'MATCH: BOT-PAT-003 at 100% (identical bot farm signature)'],
    ['Neptune Graph', 'SUCCESS', '110ms', 'BOT-FARM-001 in network of 47 accounts sharing 5 devices'],
    ['Composite Score', 'SUCCESS', '<1ms', 'Score: 0.63 | Decision: CHALLENGE'],
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Final result:').bold = True
doc.add_paragraph('• Total latency: 340ms')
doc.add_paragraph('• Risk score: 63/100')
doc.add_paragraph('• Decision: CHALLENGE')
doc.add_paragraph('• Similarity: 100% match to known bot farm listening pattern')
doc.add_paragraph('• Graph: 47-account network sharing 5 device fingerprints')
doc.add_paragraph()
doc.add_paragraph('Key indicators detected:')
doc.add_paragraph('• 661,000 streams/day (normal user: 10-200)', style='List Bullet')
doc.add_paragraph('• 0.3s average listen duration (normal: 180s)', style='List Bullet')
doc.add_paragraph('• Device shared across 47 accounts', style='List Bullet')
doc.add_paragraph('• All streaming AI-generated tracks exclusively', style='List Bullet')
doc.add_paragraph()
add_talking_point(doc, 'pgvector detects the unnatural listening pattern — no human listens to 661K tracks for 0.3 seconds each. Neptune maps the full 47-account bot farm sharing just 5 devices. Together they block fraudulent royalty payouts before the check is cut. Michael Smith operated undetected for years — this catches it in 340ms.')
doc.add_page_break()

# ===== STEP 9: DEMO 4 - IMAX =====
doc.add_heading('Step 9: Full Pipeline Demo 4 — IMAX', level=1)
doc.add_heading('"The Scalper Bot Network"', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Real-world context: ').bold = True
p.add_run('IMAX reported record $1.2B ticket sales in 2025. 51% of all web traffic is now automated (Imperva 2024 report). Executive Order 14254 (BOTS Act) specifically targets automated ticket purchasing. IMAX exploring potential sale — protecting ticket revenue is critical.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Scenario: ').bold = True
p.add_run('When a premium IMAX showing goes on sale, 200 bot sessions hit within 30 seconds. Each has superhuman interaction speed (85ms vs 3000ms human average), buys 8 tickets, and shares device fingerprints with other sessions. System detects and flags the coordinated network.')
doc.add_paragraph()
doc.add_paragraph('Navigate: Select "IMAX" → Demo page → Execute')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Input data sent:').bold = True
doc.add_paragraph()
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Parameter'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'domain'
table.rows[1].cells[1].text = 'imax'
table.rows[2].cells[0].text = 'event_type'
table.rows[2].cells[1].text = 'purchase_attempt'
table.rows[3].cells[0].text = 'entity_id'
table.rows[3].cells[1].text = 'SESS-BOT-001'
table.rows[4].cells[0].text = 'content'
table.rows[4].cells[1].text = 'Scalper bot network detected: 200 simultaneous sessions within 30 seconds targeting same premium showing. All sessions share 5 device fingerprints and 3 payment BINs.'
table.rows[5].cells[0].text = 'payload.showtime_id'
table.rows[5].cells[1].text = 'IMAX-PREM-001'
table.rows[6].cells[0].text = 'payload.quantity'
table.rows[6].cells[1].text = '8'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pipeline stages and results:').bold = True
doc.add_paragraph()
add_stage_table(doc, [
    ['Cache Check', 'MISS', '<1ms', 'First time seeing this session'],
    ['DynamoDB Ingest', 'SUCCESS', '8ms', 'Stored in multidb_poc_purchase_events'],
    ['Bedrock Embedding', 'SUCCESS', '118ms', '1024-dim vector of bot purchase behavior'],
    ['pgvector Similarity', 'SUCCESS', '150ms', 'MATCH: SCALP-002 at 100% (simultaneous sessions pattern)'],
    ['Neptune Graph', 'SUCCESS', '105ms', 'SESS-BOT-001 shares BOT-FP-0 with 4 sessions, all target IMAX-PREM-001'],
    ['Composite Score', 'SUCCESS', '<1ms', 'Score: 0.55 | Decision: FLAG'],
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Final result:').bold = True
doc.add_paragraph('• Total latency: 374ms')
doc.add_paragraph('• Risk score: 55/100')
doc.add_paragraph('• Decision: FLAG (manual review required)')
doc.add_paragraph('• Similarity: 100% match to known scalper bot pattern')
doc.add_paragraph('• Graph: 23 sessions in coordinated network, 5 shared device fingerprints')
doc.add_paragraph()
doc.add_paragraph('Key indicators detected:')
doc.add_paragraph('• Interaction speed: 85ms (human average: 3,000ms) — 35x faster than possible by hand', style='List Bullet')
doc.add_paragraph('• 200 simultaneous sessions in 30-second window', style='List Bullet')
doc.add_paragraph('• Bulk quantity: 8 tickets per session (scalper behavior)', style='List Bullet')
doc.add_paragraph('• 23 sessions sharing only 5 device fingerprints', style='List Bullet')
doc.add_paragraph('• All targeting same premium showing (Avatar 4, IMAX NYC, $35/ticket)', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Revenue protected: 23 sessions × 8 tickets × $35 = $6,440 per showing attack')
doc.add_paragraph()
add_talking_point(doc, 'Human checkout takes 3,000ms minimum — mouse movements, reading, clicking. This session completed in 85ms — physically impossible for a human. pgvector catches the behavioral anomaly. Neptune maps all 23 coordinated sessions sharing 5 device fingerprints targeting the same premium showing. Detection in 374ms, before a single ticket is sold to scalpers.')
doc.add_page_break()

# ===== STEP 10: DEMO 5 - Particle Media =====
doc.add_heading('Step 10: Full Pipeline Demo 5 — Particle Media', level=1)
doc.add_heading('"The Misinformation Campaign"', level=2)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Real-world context: ').bold = True
p.add_run('Reuters identified 40+ false AI-generated stories on Particle/NewsBreak. NBC discovered GoFundMe fraud linkage in the content. Congressional scrutiny followed. 50 million monthly users affected by the platform\'s algorithm amplifying AI-generated misinformation.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Scenario: ').bold = True
p.add_run('An AI-generated article making false health claims is published by bot account PM-BOT-001. The content signature matches known AI generation patterns. A 50-account bot network amplifies the article and links to a fraudulent GoFundMe page to monetize the fear.')
doc.add_paragraph()
doc.add_paragraph('Navigate: Select "Particle Media" → Demo page → Execute')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Input data sent:').bold = True
doc.add_paragraph()
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Parameter'
table.rows[0].cells[1].text = 'Value'
table.rows[1].cells[0].text = 'domain'
table.rows[1].cells[1].text = 'particle_media'
table.rows[2].cells[0].text = 'event_type'
table.rows[2].cells[1].text = 'content_published'
table.rows[3].cells[0].text = 'entity_id'
table.rows[3].cells[1].text = 'PM-BOT-001'
table.rows[4].cells[0].text = 'content'
table.rows[4].cells[1].text = 'BREAKING: Major pharmaceutical company admits vaccine causes severe side effects in 90 percent of recipients. Sources confirm internal documents leaked showing massive cover-up.'
table.rows[5].cells[0].text = 'payload.content_id'
table.rows[5].cells[1].text = 'MISINFO-2024-001'
table.rows[6].cells[0].text = 'payload.author_id'
table.rows[6].cells[1].text = 'BOT-AUTHOR-50'
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pipeline stages and results:').bold = True
doc.add_paragraph()
add_stage_table(doc, [
    ['Cache Check', 'MISS', '<1ms', 'New content being analyzed'],
    ['DynamoDB Ingest', 'SUCCESS', '7ms', 'Stored in multidb_poc_content_engagement_events'],
    ['Bedrock Embedding', 'SUCCESS', '128ms', '1024-dim vector of health misinfo content'],
    ['pgvector Similarity', 'SUCCESS', '132ms', 'MATCH: MISINFO-003 at 100% (vaccine misinfo + cover-up framing)'],
    ['Neptune Graph', 'SUCCESS', '108ms', 'PM-BOT-001 is MEMBER_OF amplification network (50 accounts)'],
    ['Composite Score', 'SUCCESS', '<1ms', 'Score: 0.74 | Decision: CHALLENGE'],
])
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Final result:').bold = True
doc.add_paragraph('• Total latency: 379ms')
doc.add_paragraph('• Risk score: 74/100')
doc.add_paragraph('• Decision: CHALLENGE')
doc.add_paragraph('• Similarity: 100% match to known AI-generated health misinformation pattern')
doc.add_paragraph('• Graph: Part of 50-account coordinated amplification network')
doc.add_paragraph()
doc.add_paragraph('Key indicators detected:')
doc.add_paragraph('• AI content signature: Sensationalist framing ("BREAKING", "massive cover-up")', style='List Bullet')
doc.add_paragraph('• Health misinformation: False claims about vaccine side effects', style='List Bullet')
doc.add_paragraph('• Coordinated amplification: 50 bot accounts created same week', style='List Bullet')
doc.add_paragraph('• Fraud linkage: Connected to GoFundMe disaster relief scam', style='List Bullet')
doc.add_paragraph('• Credibility score of original content: 0.12/1.00 (extremely low)', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Network structure discovered:')
doc.add_paragraph('• PM-BOT-001 → MEMBER_OF → AMPLIFICATION-NET-01', style='List Bullet')
doc.add_paragraph('• PM-BOT-001 → AMPLIFIED → MISINFO-001, MISINFO-002, MISINFO-003', style='List Bullet')
doc.add_paragraph('• AMPLIFICATION-NET-01 contains 50 member accounts', style='List Bullet')
doc.add_paragraph('• All accounts → AMPLIFIED → same 5 misinformation articles', style='List Bullet')
doc.add_paragraph()
add_talking_point(doc, 'Three independent signals converge: the content embedding detects AI-generated text signatures, the graph reveals a coordinated 50-account amplification network created the same week, and there is a fraudulent GoFundMe link designed to monetize the fear. One database sees one signal — four databases see the complete picture in 379ms.')
doc.add_page_break()

# ===== STEP 11: BRIEFING =====
doc.add_heading('Step 11: AI Investigator Briefing (Claude via Bedrock)', level=1)
doc.add_paragraph('Navigate: Click "📋 Briefing" in the sidebar')
doc.add_paragraph()
doc.add_paragraph('This generates a full investigation report using Claude Haiku 4.5 via Amazon Bedrock, synthesizing evidence from all 4 database tiers.')
doc.add_paragraph()

doc.add_heading('Example: Generate Briefing for USR-FAKE-001 (Match Group)', level=2)
doc.add_paragraph('1. Enter Entity ID: USR-FAKE-001')
doc.add_paragraph('2. Ensure domain is "Match Group"')
doc.add_paragraph('3. Click "Generate Briefing"')
doc.add_paragraph('4. Wait 3-5 seconds')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Expected output structure:').bold = True
doc.add_paragraph()
doc.add_paragraph('TITLE: "Investigator Briefing: Romance Scam Ring Detection"', style='List Bullet')
doc.add_paragraph('EXECUTIVE SUMMARY: USR-FAKE-001 identified as primary node in active scam ring with 3+ coordinated accounts sharing device DEV-SHARED-1', style='List Bullet')
doc.add_paragraph('ENTITY PROFILE: Classification: Primary Scam Account, Ring Position: Coordinator, Connected Ring Size: 7, Infrastructure: DEV-SHARED-1', style='List Bullet')
doc.add_paragraph('EVIDENCE TIMELINE:', style='List Bullet')
doc.add_paragraph('  1. Infrastructure Establishment — shared device provisioned', style='List Bullet')
doc.add_paragraph('  2. Multi-Profile Deployment — USR-FAKE-004, 007, 010 created', style='List Bullet')
doc.add_paragraph('  3. Victim Engagement — contact with Amanda (USR-REAL-003)', style='List Bullet')
doc.add_paragraph('  4. Coordinated Messaging — scripted outreach campaigns', style='List Bullet')
doc.add_paragraph('RISK ASSESSMENT: Critical (0.94) with justification', style='List Bullet')
doc.add_paragraph('RECOMMENDED ACTIONS (8 prioritized):', style='List Bullet')
doc.add_paragraph('  1. Immediate suspension — all ring accounts (2 hours)', style='List Bullet')
doc.add_paragraph('  2. Victim notification — contact Amanda (1 hour)', style='List Bullet')
doc.add_paragraph('  3. Device forensics — DEV-SHARED-1 (24 hours)', style='List Bullet')
doc.add_paragraph('  4. Law enforcement — FBI IC3 referral (24 hours)', style='List Bullet')
doc.add_paragraph('  5. Ring expansion — deeper graph analysis (48 hours)', style='List Bullet')
doc.add_paragraph('  6. Messaging analysis — extract templates (48 hours)', style='List Bullet')
doc.add_paragraph('  7. Financial tracing — payment processors (72 hours)', style='List Bullet')
doc.add_paragraph('  8. Detection rules — platform-wide deployment (1 week)', style='List Bullet')
doc.add_paragraph('CONFIDENCE SCORE: 0.92', style='List Bullet')
doc.add_paragraph()
add_talking_point(doc, 'Claude synthesizes evidence from all 4 database tiers into a production-quality investigation report. An analyst who would spend 2-4 hours writing this gets it in 5 seconds, with specific actions assigned to teams with timelines.')
doc.add_paragraph()

doc.add_heading('Other Briefings to Generate', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Entity'
table.rows[0].cells[1].text = 'Expected Focus'
table.rows[1].cells[0].text = 'J-UNKNOWN-443 (Business Wire)'
table.rows[1].cells[1].text = 'Insider trading network, unauthorized access timeline, SEC referral recommendation'
table.rows[2].cells[0].text = 'BOT-FARM-001 (UMG)'
table.rows[2].cells[1].text = 'Stream manipulation ring, royalty fraud quantification, platform demonetization recommendation'
table.rows[3].cells[0].text = 'SESS-BOT-001 (IMAX)'
table.rows[3].cells[1].text = 'Scalper network infrastructure, revenue impact ($6,440/show), BOTS Act violation'
table.rows[4].cells[0].text = 'PM-BOT-001 (Particle Media)'
table.rows[4].cells[1].text = 'AI content generation detection, amplification network, GoFundMe fraud, FTC referral'
doc.add_page_break()

# ===== STEP 12 =====
doc.add_heading('Step 12: Advanced — Custom Scenarios', level=1)
doc.add_paragraph()
doc.add_heading('Test Your Own Scam Variation', level=2)
doc.add_paragraph('Go to Semantic page, domain: Match Group, paste:')
add_code_block(doc, 'Dear beautiful lady, I am a Navy SEAL currently serving in\nAfghanistan. I saw your profile and felt an instant connection.\nMy wife died 2 years ago and I am ready to love again.\nPlease add me on Google Hangouts.')
doc.add_paragraph('Expected: 70-90% match (catches semantic variations even with different military branch, platform, and backstory)')
doc.add_paragraph()
doc.add_heading('Test a Completely New Threat', level=2)
doc.add_paragraph('Domain: Match Group, paste:')
add_code_block(doc, 'I am a cryptocurrency investor who made millions. I want to\nshare my trading secrets with you. Just send me $500 to get\nstarted and I will triple your money within 24 hours.')
doc.add_paragraph('Expected: Lower match (40-60%) — this is an investment scam, not romance scam. Different pattern category.')
doc.add_paragraph()
doc.add_heading('Verify Zero False Positives', level=2)
doc.add_paragraph('Domain: Match Group, paste legitimate messages:')
add_code_block(doc, 'Hey! Love your taste in music. I work in tech and also enjoy\nhiking on weekends. Want to grab coffee sometime?')
doc.add_paragraph('Expected: 0 matches. System correctly identifies this as safe.')
doc.add_page_break()

# ===== STEP 13: API Reference =====
doc.add_heading('Step 13: API Reference', level=1)
doc.add_paragraph('Base URL: https://nkt0mgbdn5.execute-api.us-east-1.amazonaws.com/v1')
doc.add_paragraph()

doc.add_heading('Full Pipeline', level=2)
add_code_block(doc, 'curl -X POST .../api/v1/pipeline/execute \\\n  -H "Content-Type: application/json" \\\n  -d \'{"domain":"match_group","event_type":"message_sent",\n       "content":"Your text here","payload":{"user_id":"X"},\n       "entity_id":"USR-FAKE-001"}\'')
doc.add_paragraph()
doc.add_heading('Semantic Search', level=2)
add_code_block(doc, 'curl -X POST .../api/v1/analysis/semantic/ \\\n  -d \'{"domain":"match_group","content":"Text to search",\n       "similarity_threshold":0.6,"top_k":10}\'')
doc.add_paragraph()
doc.add_heading('Graph Analysis', level=2)
add_code_block(doc, 'curl -X POST .../api/v1/analysis/graph/ \\\n  -d \'{"entity_id":"USR-FAKE-001","algorithm":"louvain","max_depth":3}\'')
doc.add_paragraph()
doc.add_heading('Claude Briefing', level=2)
add_code_block(doc, 'curl .../api/v1/briefing/USR-FAKE-001?domain=match_group')
doc.add_paragraph()
doc.add_heading('Seed Data (re-initialize)', level=2)
add_code_block(doc, 'curl -X POST .../api/v1/admin/seed/all\ncurl -X POST .../api/v1/admin/seed-embeddings')
doc.add_page_break()

# ===== APPENDIX A =====
doc.add_heading('Appendix A: Entity IDs Quick Reference', level=1)
doc.add_paragraph()
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Customer'
table.rows[0].cells[1].text = 'Entity ID'
table.rows[0].cells[2].text = 'Type'
table.rows[0].cells[3].text = 'Ring Size'
table.rows[1].cells[0].text = 'Match Group'
table.rows[1].cells[1].text = 'USR-FAKE-001'
table.rows[1].cells[2].text = 'Romance scammer'
table.rows[1].cells[3].text = '15 members, 3 devices'
table.rows[2].cells[0].text = 'Business Wire'
table.rows[2].cells[1].text = 'J-UNKNOWN-443'
table.rows[2].cells[2].text = 'Unauthorized journalist'
table.rows[2].cells[3].text = '3 actors, shared IP'
table.rows[3].cells[0].text = 'Universal Music'
table.rows[3].cells[1].text = 'BOT-FARM-001'
table.rows[3].cells[2].text = 'Stream bot account'
table.rows[3].cells[3].text = '47 accounts, 5 devices'
table.rows[4].cells[0].text = 'IMAX'
table.rows[4].cells[1].text = 'SESS-BOT-001'
table.rows[4].cells[2].text = 'Scalper bot session'
table.rows[4].cells[3].text = '23 sessions, 5 fingerprints'
table.rows[5].cells[0].text = 'Particle Media'
table.rows[5].cells[1].text = 'PM-BOT-001'
table.rows[5].cells[2].text = 'Misinfo amplifier'
table.rows[5].cells[3].text = '50 accounts, 1 network'

doc.add_paragraph()

# ===== APPENDIX B =====
doc.add_heading('Appendix B: Architecture & Decision Matrix', level=1)
doc.add_paragraph()
doc.add_heading('Score Weights per Domain', level=2)
table = doc.add_table(rows=6, cols=5)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Domain'
table.rows[0].cells[1].text = 'Graph'
table.rows[0].cells[2].text = 'Similarity'
table.rows[0].cells[3].text = 'Behavioral'
table.rows[0].cells[4].text = 'Velocity'
table.rows[1].cells[0].text = 'Business Wire'
table.rows[1].cells[1].text = '35%'
table.rows[1].cells[2].text = '30%'
table.rows[1].cells[3].text = '20%'
table.rows[1].cells[4].text = '15%'
table.rows[2].cells[0].text = 'Match Group'
table.rows[2].cells[1].text = '30%'
table.rows[2].cells[2].text = '35%'
table.rows[2].cells[3].text = '20%'
table.rows[2].cells[4].text = '15%'
table.rows[3].cells[0].text = 'UMG'
table.rows[3].cells[1].text = '25%'
table.rows[3].cells[2].text = '25%'
table.rows[3].cells[3].text = '30%'
table.rows[3].cells[4].text = '20%'
table.rows[4].cells[0].text = 'IMAX'
table.rows[4].cells[1].text = '20%'
table.rows[4].cells[2].text = '20%'
table.rows[4].cells[3].text = '25%'
table.rows[4].cells[4].text = '35%'
table.rows[5].cells[0].text = 'Particle Media'
table.rows[5].cells[1].text = '30%'
table.rows[5].cells[2].text = '35%'
table.rows[5].cells[3].text = '20%'
table.rows[5].cells[4].text = '15%'
doc.add_paragraph()

doc.add_heading('Decision Matrix', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Score Range'
table.rows[0].cells[1].text = 'Decision'
table.rows[0].cells[2].text = 'Platform Action'
table.rows[1].cells[0].text = '0.00 - 0.30'
table.rows[1].cells[1].text = 'ALLOW'
table.rows[1].cells[2].text = 'No action — normal activity'
table.rows[2].cells[0].text = '0.30 - 0.60'
table.rows[2].cells[1].text = 'FLAG'
table.rows[2].cells[2].text = 'Queue for manual review'
table.rows[3].cells[0].text = '0.60 - 0.80'
table.rows[3].cells[1].text = 'CHALLENGE'
table.rows[3].cells[2].text = 'Require identity verification or CAPTCHA'
table.rows[4].cells[0].text = '0.80 - 1.00'
table.rows[4].cells[1].text = 'BLOCK'
table.rows[4].cells[2].text = 'Immediate block, suspend account'
doc.add_paragraph()

# ===== APPENDIX C =====
doc.add_heading('Appendix C: Demo Timing Guide', level=1)
table = doc.add_table(rows=10, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Section'
table.rows[0].cells[1].text = 'Time'
table.rows[0].cells[2].text = 'Priority'
table.rows[1].cells[0].text = 'Dashboard overview + architecture'
table.rows[1].cells[1].text = '2 min'
table.rows[1].cells[2].text = 'MUST'
table.rows[2].cells[0].text = 'Demo 1: Match Group pipeline'
table.rows[2].cells[1].text = '3 min'
table.rows[2].cells[2].text = 'MUST'
table.rows[3].cells[0].text = 'Demo 2: Business Wire pipeline'
table.rows[3].cells[1].text = '2 min'
table.rows[3].cells[2].text = 'MUST'
table.rows[4].cells[0].text = 'Demo 3: UMG pipeline'
table.rows[4].cells[1].text = '2 min'
table.rows[4].cells[2].text = 'SHOULD'
table.rows[5].cells[0].text = 'Demo 4: IMAX pipeline'
table.rows[5].cells[1].text = '2 min'
table.rows[5].cells[2].text = 'SHOULD'
table.rows[6].cells[0].text = 'Demo 5: Particle Media pipeline'
table.rows[6].cells[1].text = '2 min'
table.rows[6].cells[2].text = 'SHOULD'
table.rows[7].cells[0].text = 'Claude Briefing generation'
table.rows[7].cells[1].text = '2 min'
table.rows[7].cells[2].text = 'MUST (wow factor)'
table.rows[8].cells[0].text = 'Semantic deep-dive + false positive test'
table.rows[8].cells[1].text = '2 min'
table.rows[8].cells[2].text = 'If time allows'
table.rows[9].cells[0].text = 'TOTAL'
table.rows[9].cells[1].text = '~17 min'
table.rows[9].cells[2].text = ''

# Save
output = '/Users/haliasgh/DMS_local_converter/multi-db-poc/docs/Multi-DB_AI_POC_Demo_Walkthrough.docx'
doc.save(output)
print(f'Generated: {output}')
